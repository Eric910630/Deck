"""
测试LLM理解Demo文档.docx
使用新的简化prompt（只有背景和方向，没有强制限定）
"""

import asyncio
from pathlib import Path
from loguru import logger
import json

from llm_service import LLMService
from human_centered_analyzer import HumanCenteredAnalyzer
from enhanced_ppt_parser import EnhancedPPTParser


def extract_docx_content(docx_path: str) -> str:
    """从docx文件中提取文本内容"""
    try:
        from docx import Document
        
        doc = Document(docx_path)
        paragraphs = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        
        return "\n".join(paragraphs)
    except ImportError:
        logger.error("需要安装python-docx: pip install python-docx")
        return ""
    except Exception as e:
        logger.error(f"读取docx文件失败: {e}")
        return ""


async def test_llm_understand_demo_docx():
    """测试LLM理解Demo文档.docx"""
    logger.info("="*80)
    logger.info("测试：LLM理解Demo文档.docx（使用新的简化prompt）")
    logger.info("="*80)
    
    # 1. 读取Demo文档.docx
    docx_path = Path("Demo文档.docx")
    if not docx_path.exists():
        logger.error(f"文件不存在: {docx_path}")
        return
    
    logger.info(f"📄 读取文件: {docx_path}")
    docx_content = extract_docx_content(str(docx_path))
    logger.info(f"   文档长度: {len(docx_content)}字符")
    logger.info(f"   文档预览: {docx_content[:500]}...")
    
    # 2. 创建结构数据（用于HumanCenteredAnalyzer）
    # 使用一个简单的框架PPT结构作为基础
    framework_path = Path("demo_filled.pptx")
    if framework_path.exists():
        parser = EnhancedPPTParser(str(framework_path))
        framework_structure = parser.extract_structure_enhanced()
    else:
        # 如果没有框架PPT，创建一个基本结构
        framework_structure = {
            "slide_count": 1,
            "slide_width": 33.867,
            "slide_height": 19.05,
            "slides": [{
                "slide_index": 0,
                "shapes": [],
                "placeholders": []
            }]
        }
    
    # 3. 创建docx结构（模拟PPT结构，但使用docx内容）
    docx_structure = {
        "slide_count": 1,
        "slide_width": 33.867,
        "slide_height": 19.05,
        "slides": [{
            "slide_index": 0,
            "shapes": [{
                "shape_id": i,
                "text": para,
                "format": {}
            } for i, para in enumerate(docx_content.split('\n') if docx_content else [])],
            "placeholders": []
        }]
    }
    
    # 4. 初始化LLM服务和HumanCenteredAnalyzer
    logger.info("🤖 初始化LLM服务...")
    llm_service = LLMService()
    
    logger.info("📊 初始化HumanCenteredAnalyzer...")
    analyzer = HumanCenteredAnalyzer(
        structure_data=docx_structure,
        raw_text=docx_content,
        llm_service=llm_service
    )
    
    # 5. 执行分析
    logger.info("🔍 开始LLM理解分析...")
    logger.info("   使用新的简化prompt：")
    logger.info("   - 背景：中国职场的述职汇报专家，专门为副总裁级别的职场高管筹备述职内容")
    logger.info("   - 方向：先通读了解核心思想 → 再细分板块拆解 → 然后深入探寻每个板块")
    logger.info("")
    
    try:
        human_analysis = await analyzer.analyze_all()
        
        # 6. 输出分析结果到文档
        output_path = Path("LLM_理解_Demo文档_分析结果.md")
        logger.info(f"📝 输出分析结果到: {output_path}")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# LLM理解Demo文档.docx - 分析结果\n\n")
            f.write("## 📋 说明\n\n")
            f.write("本文档是使用新的简化prompt（只有背景和方向，没有强制限定）对Demo文档.docx进行理解的结果。\n\n")
            f.write("**Prompt特点**：\n")
            f.write("- ✅ 背景：中国职场的述职汇报专家，专门为副总裁级别的职场高管筹备述职内容\n")
            f.write("- ✅ 方向：先通读了解核心思想 → 再细分板块拆解 → 然后深入探寻每个板块\n")
            f.write("- ✅ 没有强制限定和格式要求，让LLM自由理解\n\n")
            f.write("---\n\n")
            
            # 第1层：通读理解
            f.write("## 第1层：通读理解\n\n")
            layer1 = human_analysis.get("layer_1_overall_understanding", {}).get("data", {})
            f.write(f"### 核心主题\n\n{layer1.get('core_theme', '未识别')}\n\n")
            f.write(f"### 核心思想\n\n{layer1.get('core_idea', '未识别')}\n\n")
            f.write(f"### 文档目的\n\n{layer1.get('purpose', '未识别')}\n\n")
            f.write(f"### 目标受众\n\n{layer1.get('target_audience', '未识别')}\n\n")
            f.write(f"### 核心价值主张\n\n")
            for i, vp in enumerate(layer1.get('value_propositions', []), 1):
                f.write(f"{i}. {vp}\n")
            f.write("\n")
            f.write(f"### 关键短语\n\n")
            for i, phrase in enumerate(layer1.get('key_phrases', [])[:10], 1):
                f.write(f"{i}. {phrase}\n")
            f.write("\n")
            f.write(f"### 统计信息\n\n")
            f.write(f"- 文档总长度: {layer1.get('text_length', 0)}字符\n")
            f.write(f"- 幻灯片数: {layer1.get('total_slides', 0)}\n\n")
            f.write("---\n\n")
            
            # 第2层：板块拆分
            f.write("## 第2层：板块拆分\n\n")
            layer2 = human_analysis.get("layer_2_sections", {}).get("data", {})
            f.write(f"### 板块总数\n\n{layer2.get('total_sections', 0)}个板块\n\n")
            f.write(f"### 各板块详情\n\n")
            for section in layer2.get('sections', []):
                f.write(f"#### 板块{section.get('section_index', 0)}: {section.get('theme', '未命名')}\n\n")
                f.write(f"**核心思想**: {section.get('core_idea', '未识别')}\n\n")
                f.write(f"**内容摘要**: {section.get('content_summary', '无')}\n\n")
                f.write(f"**涉及幻灯片**: {', '.join(map(str, section.get('slides', [])))}\n\n")
            f.write("---\n\n")
            
            # 【新增】润色结果
            f.write("## 润色结果\n\n")
            f.write("### 说明\n\n")
            f.write("以下是对各板块内容进行PPT展示层面润色的结果，将文档内容润色成适合PPT展示的文案。\n\n")
            
            # 为每个板块进行润色（这里需要调用ContentPolisher）
            from content_polisher import ContentPolisher
            from presentation_planner import PresentationPlanner
            from layout_planner import LayoutPlanner
            
            content_polisher = ContentPolisher(llm_service)
            presentation_planner = PresentationPlanner(llm_service)
            layout_planner = LayoutPlanner(llm_service)
            
            for section in layer2.get('sections', []):
                section_idx = section.get('section_index', 0)
                section_analysis = {
                    "theme": section.get('theme', ''),
                    "core_idea": section.get('core_idea', ''),
                    "content_summary": section.get('content_summary', '')
                }
                
                # 获取论证信息
                layer3 = human_analysis.get("layer_3_arguments", {}).get("data", {})
                arguments = layer3.get('arguments', [])
                if section_idx < len(arguments):
                    arg = arguments[section_idx]
                    section_analysis.update({
                        "core_content": arg.get('core_content', ''),
                        "specific_arguments": arg.get('specific_arguments', []),
                        "core_evidence": arg.get('core_evidence', []),
                        "data_points": arg.get('data_points', [])
                    })
                
                f.write(f"#### 板块{section_idx}: {section.get('theme', '未命名')}\n\n")
                
                try:
                    # 润色
                    polished_slides = await content_polisher.polish_section(
                        section_analysis=section_analysis,
                        section_index=section_idx
                    )
                    
                    f.write(f"**润色后的幻灯片**（共{len(polished_slides)}张）:\n\n")
                    for slide in polished_slides:
                        f.write(f"**幻灯片{slide.get('slide_index', 0)}**: {slide.get('title', '')}\n")
                        f.write(f"- 内容: {slide.get('content', '')}\n")
                        f.write(f"- 内容类型: {slide.get('content_type', '')}\n")
                        visual = slide.get('visual_elements', {})
                        if visual.get('needs_table') or visual.get('needs_chart') or visual.get('needs_cards'):
                            f.write(f"- 视觉元素: {visual.get('notes', '需要视觉元素')}\n")
                        # 详细展开视觉元素
                        visual_detail = slide.get('visual_elements_detail', [])
                        if visual_detail:
                            f.write(f"- 视觉元素详细展开（共{len(visual_detail)}个元素）:\n")
                            for elem in visual_detail:
                                element_id = elem.get('element_id', f"{elem.get('element_type', 'unknown')}_{elem.get('element_index', 0)}")
                                f.write(f"  * 元素{elem.get('element_index', 0)} (ID: {element_id}, 类型: {elem.get('element_type', 'unknown')}): {elem.get('title', '无标题')}\n")
                                if elem.get('content'):
                                    f.write(f"    - 内容: {elem.get('content', '')}\n")
                                if elem.get('data'):
                                    f.write(f"    - 数据: {elem.get('data', '')}\n")
                                if elem.get('description'):
                                    f.write(f"    - 说明: {elem.get('description', '')}\n")
                        f.write("\n")
                    
                    # 展示策划
                    presentation_plan = await presentation_planner.plan_presentation(
                        polished_slides=polished_slides,
                        section_theme=section.get('theme', '')
                    )
                    
                    f.write(f"**展示策划**（共{len(presentation_plan)}张）:\n\n")
                    for plan in presentation_plan:
                        f.write(f"**幻灯片{plan.get('slide_index', 0)}**:\n")
                        f.write(f"- 布局类型: {plan.get('layout_type', '')}\n")
                        f.write(f"- 布局描述: {plan.get('layout_description', '')}\n")
                        guidance = plan.get('visual_guidance', {})
                        if guidance:
                            f.write(f"- 视觉指导:\n")
                            f.write(f"  - 字体大小: {guidance.get('font_size', '')}\n")
                            f.write(f"  - 字体粗细: {guidance.get('font_weight', '')}\n")
                            f.write(f"  - 对齐方式: {guidance.get('alignment', '')}\n")
                            f.write(f"  - 间距: {guidance.get('spacing', '')}\n")
                            f.write(f"  - 配色: {guidance.get('color_scheme', '')}\n")
                            if guidance.get('other_notes'):
                                f.write(f"  - 其他说明: {guidance.get('other_notes', '')}\n")
                        f.write("\n")
                    
                    # 布局规划
                    layout_plans = await layout_planner.plan_layout(
                        polished_slides=polished_slides,
                        presentation_plan=presentation_plan
                    )
                    
                    f.write("*********** 布局规划（新增） ***********\n\n")
                    f.write(f"**布局规划**（共{len(layout_plans)}张）:\n\n")
                    for layout_plan in layout_plans:
                        slide_idx = layout_plan.get('slide_index', 0)
                        plan_data = layout_plan.get('layout_plan', {})
                        f.write(f"**幻灯片{slide_idx}**:\n")
                        f.write(f"- 整体布局结构: {plan_data.get('overall_structure', '')}\n\n")
                        
                        # 元素位置
                        element_positions = plan_data.get('element_positions', [])
                        if element_positions:
                            f.write(f"- 元素位置（共{len(element_positions)}个元素）:\n")
                            for elem_pos in element_positions:
                                f.write(f"  * {elem_pos.get('element_id', '')} ({elem_pos.get('element_type', '')}):\n")
                                f.write(f"    - 位置: {elem_pos.get('position_description', '')}\n")
                                f.write(f"    - 尺寸: {elem_pos.get('size_description', '')}\n")
                                f.write(f"    - 对齐: {elem_pos.get('alignment', '')}\n")
                                spacing = elem_pos.get('spacing', {})
                                if spacing:
                                    f.write(f"    - 间距: 上{elem_pos.get('spacing', {}).get('margin_top', '')}, 下{elem_pos.get('spacing', {}).get('margin_bottom', '')}, 左{elem_pos.get('spacing', {}).get('margin_left', '')}, 右{elem_pos.get('spacing', {}).get('margin_right', '')}\n")
                            f.write("\n")
                        
                        # 元素间距
                        element_spacing = plan_data.get('element_spacing', {})
                        if element_spacing:
                            f.write(f"- 元素间距:\n")
                            f.write(f"  - 元素之间: {element_spacing.get('between_elements', '')}\n")
                            f.write(f"  - 内边距: {element_spacing.get('internal_padding', '')}\n\n")
                        
                        # 视觉层次
                        visual_hierarchy = plan_data.get('visual_hierarchy', '')
                        if visual_hierarchy:
                            f.write(f"- 视觉层次: {visual_hierarchy}\n\n")
                        
                        # 设计规范
                        design_specs = plan_data.get('design_specifications', '')
                        if design_specs:
                            f.write(f"- 设计规范: {design_specs}\n\n")
                    
                    f.write("*********** 布局规划结束 ***********\n\n")
                    
                except Exception as e:
                    logger.error(f"板块{section_idx}润色/策划失败: {e}", exc_info=True)
                    f.write(f"*润色/策划失败: {e}*\n\n")
            
            f.write("---\n\n")
            
            # 第3层：论证逻辑
            f.write("## 第3层：论证逻辑\n\n")
            layer3 = human_analysis.get("layer_3_arguments", {}).get("data", {})
            f.write(f"### 有论证的板块数\n\n{layer3.get('total_sections_with_arguments', 0)}个板块\n\n")
            f.write(f"### 各板块论证详情\n\n")
            for arg in layer3.get('arguments', []):
                f.write(f"#### 板块{arg.get('section_index', 0)}: {arg.get('section_theme', '未命名')}\n\n")
                f.write(f"**核心内容**: {arg.get('core_content', '未识别')}\n\n")
                f.write(f"**核心思想**: {arg.get('core_idea', '未识别')}\n\n")
                f.write(f"**具体论点**:\n")
                for i, point in enumerate(arg.get('specific_arguments', []), 1):
                    f.write(f"{i}. {point}\n")
                f.write("\n")
                f.write(f"**核心论据**:\n")
                for i, evidence in enumerate(arg.get('core_evidence', []), 1):
                    f.write(f"{i}. {evidence}\n")
                f.write("\n")
                f.write(f"**数据点**:\n")
                for i, data in enumerate(arg.get('data_points', []), 1):
                    f.write(f"{i}. {data}\n")
                f.write("\n")
                f.write(f"**论证类型**: {', '.join(arg.get('argument_types', []))}\n\n")
            f.write("---\n\n")
            
            # 第4层：支撑材料
            f.write("## 第4层：支撑材料\n\n")
            layer4 = human_analysis.get("layer_4_supporting_materials", {}).get("data", {})
            f.write(f"### 数据点总数\n\n{layer4.get('total_data_points', 0)}个\n\n")
            f.write(f"### 案例总数\n\n{layer4.get('total_cases', 0)}个\n\n")
            materials = layer4.get('materials', {})
            if materials.get('data_points'):
                f.write(f"### 数据点详情\n\n")
                for i, dp in enumerate(materials.get('data_points', [])[:10], 1):
                    f.write(f"{i}. {dp.get('data', '')} (上下文: {dp.get('context', '')})\n")
                f.write("\n")
            if materials.get('cases'):
                f.write(f"### 案例详情\n\n")
                for i, case in enumerate(materials.get('cases', [])[:10], 1):
                    f.write(f"{i}. {case.get('content', '')}\n")
                f.write("\n")
            f.write("---\n\n")
            
            # 第5层：表达风格
            f.write("## 第5层：表达风格\n\n")
            layer5 = human_analysis.get("layer_5_expression_style", {}).get("data", {})
            f.write(f"### 正式程度\n\n{layer5.get('formality_level', '未识别')}\n\n")
            f.write(f"### 语调\n\n{layer5.get('tone', '未识别')}\n\n")
            f.write(f"### 文化特征\n\n")
            for feature in layer5.get('cultural_features', []):
                f.write(f"- {feature}\n")
            f.write("\n")
            f.write(f"### 数字使用\n\n{layer5.get('use_of_numbers', 0)}次\n\n")
            f.write(f"### 表情符号使用\n\n{layer5.get('use_of_emojis', 0)}次\n\n")
            f.write("---\n\n")
            
            # 第6层：呈现形式
            f.write("## 第6层：呈现形式\n\n")
            layer6 = human_analysis.get("layer_6_presentation_form", {}).get("data", {})
            layout = layer6.get('layout_style', {})
            f.write(f"### 布局风格\n\n")
            f.write(f"- 宽高比: {layout.get('aspect_ratio', '未识别')}\n")
            f.write(f"- 宽度: {layout.get('width_cm', 0)}cm\n")
            f.write(f"- 高度: {layout.get('height_cm', 0)}cm\n\n")
            typography = layer6.get('typography', {})
            f.write(f"### 字体排版\n\n")
            f.write(f"- 字体大小: {', '.join(map(str, typography.get('font_sizes', [])))}\n")
            f.write(f"- 字体名称: {', '.join(typography.get('font_names', []))}\n")
            f.write(f"- 加粗使用: {typography.get('bold_usage_count', 0)}次\n\n")
            f.write("---\n\n")
            
            # 原始JSON（用于调试）
            f.write("## 原始JSON数据（用于调试）\n\n")
            f.write("```json\n")
            f.write(json.dumps(human_analysis, indent=2, ensure_ascii=False))
            f.write("\n```\n")
        
        logger.info(f"✅ 分析结果已保存到: {output_path}")
        logger.info(f"   文件大小: {output_path.stat().st_size}字节")
        
    except Exception as e:
        logger.error(f"❌ 分析失败: {e}", exc_info=True)


if __name__ == "__main__":
    asyncio.run(test_llm_understand_demo_docx())

