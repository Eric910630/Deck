#!/usr/bin/env python3
"""
深度分析Demo文档.docx - 通过内容语义和格式识别层次结构
"""

import json
from docx import Document
from pathlib import Path
import re

def deep_analyze_docx(docx_path: str):
    """深度分析docx文档，识别所有可能的层次"""
    
    doc = Document(docx_path)
    
    analysis = {
        "file_info": {
            "path": docx_path,
            "size_bytes": Path(docx_path).stat().st_size
        },
        "layers": {}
    }
    
    # ========== 层次1: 物理结构层 ==========
    physical_structure = {
        "total_paragraphs": len(doc.paragraphs),
        "total_tables": len(doc.tables),
        "total_sections": len(doc.sections),
        "paragraphs_by_style": {},
        "paragraphs_by_format": {
            "bold_paragraphs": [],
            "large_font_paragraphs": [],
            "indented_paragraphs": []
        }
    }
    
    for para_idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "Unknown"
        if style_name not in physical_structure["paragraphs_by_style"]:
            physical_structure["paragraphs_by_style"][style_name] = []
        physical_structure["paragraphs_by_style"][style_name].append({
            "text": para.text[:100],
            "index": para_idx
        })
        
        # 检查格式特征
        max_font_size = 0
        has_bold = False
        for run in para.runs:
            if run.font.size and run.font.size.pt:
                max_font_size = max(max_font_size, run.font.size.pt)
            if run.bold:
                has_bold = True
        
        if has_bold:
            physical_structure["paragraphs_by_format"]["bold_paragraphs"].append({
                "text": para.text[:100],
                "index": para_idx
            })
        
        if max_font_size >= 20:
            physical_structure["paragraphs_by_format"]["large_font_paragraphs"].append({
                "text": para.text[:100],
                "font_size": max_font_size,
                "index": para_idx
            })
    
    analysis["layers"]["layer_1_physical"] = {
        "name": "物理结构层",
        "description": "文档的物理结构：段落、表格、章节等",
        "data": physical_structure
    }
    
    # ========== 层次2: 格式特征层 ==========
    format_features = {
        "font_sizes": set(),
        "font_names": set(),
        "bold_runs": [],
        "italic_runs": [],
        "colored_runs": [],
        "paragraph_formats": []
    }
    
    for para in doc.paragraphs:
        para_format = {
            "text": para.text[:50],
            "alignment": str(para.alignment) if para.alignment else None,
            "left_indent": para.paragraph_format.left_indent.pt if para.paragraph_format.left_indent else 0,
            "first_line_indent": para.paragraph_format.first_line_indent.pt if para.paragraph_format.first_line_indent else 0,
            "space_before": para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0,
            "space_after": para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0
        }
        format_features["paragraph_formats"].append(para_format)
        
        for run in para.runs:
            if run.font.size and run.font.size.pt:
                format_features["font_sizes"].add(run.font.size.pt)
            if run.font.name:
                format_features["font_names"].add(run.font.name)
            if run.bold and run.text.strip():
                format_features["bold_runs"].append({
                    "text": run.text[:50],
                    "font_size": run.font.size.pt if run.font.size and run.font.size.pt else None
                })
            if run.italic and run.text.strip():
                format_features["italic_runs"].append(run.text[:50])
            if run.font.color and run.font.color.rgb:
                format_features["colored_runs"].append(run.text[:50])
    
    format_features["font_sizes"] = sorted(list(format_features["font_sizes"]))
    format_features["font_names"] = list(format_features["font_names"])
    
    analysis["layers"]["layer_2_format"] = {
        "name": "格式特征层",
        "description": "文本格式特征：字体、字号、加粗、颜色、缩进等",
        "data": format_features
    }
    
    # ========== 层次3: 内容语义层 ==========
    semantic_blocks = []
    current_block = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 判断是否为标题（通过格式特征）
        is_heading = False
        heading_level = 0
        
        # 检查1: 字体大小（大字体可能是标题）
        max_font_size = 0
        has_bold = False
        for run in para.runs:
            if run.font.size and run.font.size.pt:
                max_font_size = max(max_font_size, run.font.size.pt)
            if run.bold:
                has_bold = True
        
        # 检查2: 文本长度（短文本可能是标题）
        is_short = len(text) < 50
        
        # 检查3: 是否包含数字编号（如"1. "、"一、"等）
        has_numbering = bool(re.match(r'^[\d一二三四五六七八九十]+[\.、]', text))
        
        # 检查4: 是否全为加粗
        all_bold = all(run.bold for run in para.runs if run.text.strip())
        
        # 综合判断
        if (max_font_size >= 20 and has_bold) or (is_short and all_bold) or has_numbering:
            is_heading = True
            if max_font_size >= 24:
                heading_level = 1
            elif max_font_size >= 18:
                heading_level = 2
            else:
                heading_level = 3
            
            # 保存之前的块
            if current_block:
                semantic_blocks.append(current_block)
            
            # 开始新块
            current_block = {
                "heading": text,
                "heading_level": heading_level,
                "heading_format": {
                    "font_size": max_font_size,
                    "is_bold": has_bold,
                    "has_numbering": has_numbering
                },
                "content": []
            }
        else:
            # 添加到当前块的内容
            if current_block:
                current_block["content"].append({
                    "text": text,
                    "font_size": max_font_size,
                    "is_bold": has_bold
                })
            else:
                # 如果没有标题，创建一个匿名块
                current_block = {
                    "heading": None,
                    "heading_level": 0,
                    "content": [{
                        "text": text,
                        "font_size": max_font_size,
                        "is_bold": has_bold
                    }]
                }
    
    # 保存最后一个块
    if current_block:
        semantic_blocks.append(current_block)
    
    analysis["layers"]["layer_3_semantic"] = {
        "name": "内容语义层",
        "description": "基于内容语义和格式特征识别的标题-内容块结构",
        "data": {
            "total_blocks": len(semantic_blocks),
            "blocks": semantic_blocks
        }
    }
    
    # ========== 层次4: 列表结构层 ==========
    list_structure = {
        "numbered_lists": [],
        "bullet_lists": [],
        "indented_items": []
    }
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检查编号列表
        numbered_match = re.match(r'^[\d一二三四五六七八九十]+[\.、]\s*(.+)', text)
        if numbered_match:
            list_structure["numbered_lists"].append({
                "number": numbered_match.group(1) if numbered_match.group(1) else text[0],
                "content": numbered_match.group(2) if len(numbered_match.groups()) > 1 else text,
                "full_text": text
            })
        
        # 检查项目符号
        bullet_match = re.match(r'^[•·▪▫○●■□]\s*(.+)', text)
        if bullet_match:
            list_structure["bullet_lists"].append({
                "bullet": text[0],
                "content": bullet_match.group(1),
                "full_text": text
            })
        
        # 检查缩进（可能是列表）
        if para.paragraph_format.left_indent and para.paragraph_format.left_indent.pt > 0:
            list_structure["indented_items"].append({
                "text": text[:100],
                "indent_pt": para.paragraph_format.left_indent.pt
            })
    
    analysis["layers"]["layer_4_lists"] = {
        "name": "列表结构层",
        "description": "文档中的列表结构：编号列表、项目符号、缩进列表",
        "data": list_structure
    }
    
    # ========== 层次5: 表格数据层 ==========
    table_analysis = []
    for i, table in enumerate(doc.tables):
        table_info = {
            "index": i,
            "rows": len(table.rows),
            "columns": len(table.columns) if table.rows else 0,
            "structure": [],
            "header_row": None,
            "data_rows": []
        }
        
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            table_info["structure"].append(row_data)
            
            # 第一行通常是表头
            if row_idx == 0:
                table_info["header_row"] = row_data
            else:
                table_info["data_rows"].append({
                    "row_index": row_idx,
                    "data": row_data
                })
        
        table_analysis.append(table_info)
    
    analysis["layers"]["layer_5_tables"] = {
        "name": "表格数据层",
        "description": "文档中的表格结构和数据",
        "data": table_analysis
    }
    
    # ========== 层次6: 主题/话题层 ==========
    # 通过关键词和内容聚类识别主题
    topics = {}
    keywords_patterns = {
        "业务相关": ["业务", "销售", "客户", "市场", "产品"],
        "技术相关": ["技术", "系统", "平台", "开发", "实现"],
        "管理相关": ["管理", "流程", "规范", "制度", "标准"],
        "数据相关": ["数据", "分析", "统计", "报表", "指标"]
    }
    
    for block in semantic_blocks:
        block_text = block["heading"] or ""
        if block["content"]:
            block_text += " " + " ".join([c["text"] for c in block["content"][:3]])
        
        for topic, keywords in keywords_patterns.items():
            if any(keyword in block_text for keyword in keywords):
                if topic not in topics:
                    topics[topic] = []
                topics[topic].append({
                    "heading": block["heading"],
                    "content_preview": [c["text"][:50] for c in block["content"][:2]]
                })
                break
    
    analysis["layers"]["layer_6_topics"] = {
        "name": "主题/话题层",
        "description": "基于关键词识别的内容主题分类",
        "data": topics
    }
    
    # ========== 层次7: 逻辑关系层 ==========
    # 识别内容之间的逻辑关系
    logical_relations = {
        "sequential": [],  # 顺序关系
        "hierarchical": [],  # 层级关系
        "comparative": []  # 对比关系
    }
    
    for i, block in enumerate(semantic_blocks):
        if i > 0:
            prev_block = semantic_blocks[i-1]
            # 顺序关系：连续的块
            logical_relations["sequential"].append({
                "from": prev_block["heading"] or f"块{i}",
                "to": block["heading"] or f"块{i+1}"
            })
        
        # 层级关系：标题级别
        if block["heading_level"] > 0:
            logical_relations["hierarchical"].append({
                "heading": block["heading"],
                "level": block["heading_level"],
                "sub_items": len(block["content"])
            })
    
    # 检查表格中的对比关系
    for table_info in table_analysis:
        if table_info["header_row"] and len(table_info["header_row"]) >= 2:
            logical_relations["comparative"].append({
                "type": "table_comparison",
                "table_index": table_info["index"],
                "comparison_dimensions": table_info["header_row"]
            })
    
    analysis["layers"]["layer_7_logic"] = {
        "name": "逻辑关系层",
        "description": "内容之间的逻辑关系：顺序、层级、对比等",
        "data": logical_relations
    }
    
    return analysis

if __name__ == "__main__":
    docx_path = "Demo文档.docx"
    
    print("="*80)
    print("Demo文档.docx 深度层次拆解")
    print("="*80)
    print(f"\n正在深度分析: {docx_path}\n")
    
    analysis = deep_analyze_docx(docx_path)
    
    # 输出每个层次
    for layer_key, layer_data in analysis["layers"].items():
        print("\n" + "="*80)
        print(f"【{layer_data['name']}】")
        print("="*80)
        print(f"描述: {layer_data['description']}\n")
        
        if layer_key == "layer_1_physical":
            data = layer_data["data"]
            print(f"总段落数: {data['total_paragraphs']}")
            print(f"总表格数: {data['total_tables']}")
            print(f"使用的样式: {list(data['paragraphs_by_style'].keys())}")
            print(f"大字体段落数: {len(data['paragraphs_by_format']['large_font_paragraphs'])}")
            print(f"加粗段落数: {len(data['paragraphs_by_format']['bold_paragraphs'])}")
        
        elif layer_key == "layer_2_format":
            data = layer_data["data"]
            print(f"使用的字体: {data['font_names']}")
            print(f"使用的字号: {sorted(data['font_sizes'])}")
            print(f"加粗文本段数: {len(data['bold_runs'])}")
            print(f"缩进段落数: {len([p for p in data['paragraph_formats'] if p['left_indent'] > 0])}")
        
        elif layer_key == "layer_3_semantic":
            data = layer_data["data"]
            print(f"识别的内容块数: {data['total_blocks']}")
            print(f"\n内容块详情:")
            for i, block in enumerate(data["blocks"][:10], 1):
                heading = block["heading"] or "(无标题)"
                print(f"\n  块 {i}: {heading}")
                print(f"    级别: {block['heading_level']}")
                print(f"    内容段落数: {len(block['content'])}")
                if block["content"]:
                    preview = block["content"][0]["text"][:60]
                    print(f"    内容预览: {preview}...")
        
        elif layer_key == "layer_4_lists":
            data = layer_data["data"]
            print(f"编号列表项: {len(data['numbered_lists'])}")
            print(f"项目符号项: {len(data['bullet_lists'])}")
            print(f"缩进项: {len(data['indented_items'])}")
            if data['numbered_lists']:
                print(f"\n编号列表示例:")
                for item in data['numbered_lists'][:3]:
                    print(f"  {item['full_text'][:60]}...")
        
        elif layer_key == "layer_5_tables":
            data = layer_data["data"]
            print(f"表格数量: {len(data)}")
            for table in data:
                print(f"\n表格 {table['index'] + 1}:")
                print(f"  行数: {table['rows']}, 列数: {table['columns']}")
                if table['header_row']:
                    print(f"  表头: {table['header_row']}")
                if table['data_rows']:
                    print(f"  数据行示例:")
                    for row in table['data_rows'][:2]:
                        print(f"    {row['data']}")
        
        elif layer_key == "layer_6_topics":
            data = layer_data["data"]
            print(f"识别的话题数: {len(data)}")
            for topic, blocks in data.items():
                print(f"\n话题: {topic}")
                print(f"  相关块数: {len(blocks)}")
                for block in blocks[:2]:
                    print(f"    - {block['heading']}")
        
        elif layer_key == "layer_7_logic":
            data = layer_data["data"]
            print(f"顺序关系: {len(data['sequential'])}")
            print(f"层级关系: {len(data['hierarchical'])}")
            print(f"对比关系: {len(data['comparative'])}")
            if data['hierarchical']:
                print(f"\n层级结构:")
                for rel in data['hierarchical'][:5]:
                    print(f"  {rel['heading']} (级别{rel['level']}, {rel['sub_items']}个子项)")
    
    # 保存完整分析
    output_file = "demo_docx_deep_analysis.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(analysis, f, indent=2, ensure_ascii=False)
    
    print("\n\n" + "="*80)
    print("分析完成")
    print("="*80)
    print(f"✅ 完整分析结果已保存到: {output_file}")
    print(f"   识别层次数: {len(analysis['layers'])}")
    print(f"   内容块数: {analysis['layers']['layer_3_semantic']['data']['total_blocks']}")

