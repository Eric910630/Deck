#!/usr/bin/env python3
"""
手动分析Demo文档.docx的结构和层次
不使用系统，直接由AI进行深度解读
"""

import json
from docx import Document
from pathlib import Path

def analyze_docx_structure(docx_path: str):
    """深度分析docx文档的结构和层次"""
    
    doc = Document(docx_path)
    
    analysis = {
        "file_info": {
            "path": docx_path,
            "size_bytes": Path(docx_path).stat().st_size,
            "paragraph_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
            "sections_count": len(doc.sections)
        },
        "document_structure": {
            "sections": [],
            "paragraphs": [],
            "tables": [],
            "styles_used": set(),
            "hierarchical_structure": []
        },
        "content_analysis": {
            "headings": [],
            "body_text": [],
            "lists": [],
            "tables": [],
            "metadata": {}
        },
        "semantic_layers": []
    }
    
    # 分析段落
    current_heading_level = 0
    current_section = None
    paragraph_stack = []
    
    for i, para in enumerate(doc.paragraphs):
        para_info = {
            "index": i,
            "text": para.text.strip(),
            "style": para.style.name if para.style else "Normal",
            "runs": [],
            "level": None,
            "is_heading": False,
            "is_list": False,
            "is_empty": not para.text.strip()
        }
        
        # 分析样式
        if para.style:
            analysis["document_structure"]["styles_used"].add(para.style.name)
        
        # 判断是否为标题
        if para.style and para.style.name.startswith('Heading'):
            para_info["is_heading"] = True
            try:
                para_info["level"] = int(para.style.name.replace('Heading ', ''))
            except:
                para_info["level"] = 1
        
        # 分析文本运行（runs）
        for run in para.runs:
            run_info = {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name if run.font.name else None,
                "font_size": run.font.size.pt if run.font.size else None,
                "font_color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None
            }
            para_info["runs"].append(run_info)
        
        # 判断是否为列表
        if para.style and ('List' in para.style.name or para.style.name.startswith('List')):
            para_info["is_list"] = True
        
        # 检查段落格式中的列表信息
        if para.paragraph_format.left_indent or para.paragraph_format.first_line_indent:
            para_info["is_list"] = True
        
        analysis["document_structure"]["paragraphs"].append(para_info)
        
        # 构建层次结构
        if para_info["is_heading"]:
            # 遇到新标题，处理之前的段落栈
            if paragraph_stack:
                analysis["document_structure"]["hierarchical_structure"].append({
                    "heading": paragraph_stack[0] if paragraph_stack else None,
                    "content": paragraph_stack[1:] if len(paragraph_stack) > 1 else []
                })
                paragraph_stack = []
            
            paragraph_stack.append(para_info)
            analysis["content_analysis"]["headings"].append(para_info)
        elif para_info["text"]:
            paragraph_stack.append(para_info)
            if not para_info["is_heading"]:
                analysis["content_analysis"]["body_text"].append(para_info)
    
    # 处理最后一个段落栈
    if paragraph_stack:
        analysis["document_structure"]["hierarchical_structure"].append({
            "heading": paragraph_stack[0] if paragraph_stack and paragraph_stack[0]["is_heading"] else None,
            "content": paragraph_stack[1:] if len(paragraph_stack) > 1 else paragraph_stack
        })
    
    # 分析表格
    for i, table in enumerate(doc.tables):
        table_info = {
            "index": i,
            "rows": len(table.rows),
            "columns": len(table.columns) if table.rows else 0,
            "cells": [],
            "structure": []
        }
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_data.append(cell_text)
                table_info["cells"].append({
                    "row": row_idx,
                    "column": cell_idx,
                    "text": cell_text
                })
            table_info["structure"].append(row_data)
        
        analysis["document_structure"]["tables"].append(table_info)
        analysis["content_analysis"]["tables"].append(table_info)
    
    # 分析章节
    for i, section in enumerate(doc.sections):
        section_info = {
            "index": i,
            "page_width": section.page_width.inches if section.page_width else None,
            "page_height": section.page_height.inches if section.page_height else None,
            "margin_left": section.left_margin.inches if section.left_margin else None,
            "margin_right": section.right_margin.inches if section.right_margin else None,
            "margin_top": section.top_margin.inches if section.top_margin else None,
            "margin_bottom": section.bottom_margin.inches if section.bottom_margin else None
        }
        analysis["document_structure"]["sections"].append(section_info)
    
    # 转换为可序列化的格式
    analysis["document_structure"]["styles_used"] = list(analysis["document_structure"]["styles_used"])
    
    return analysis

def identify_semantic_layers(analysis):
    """识别语义层次"""
    
    layers = []
    
    # 层次1: 文档元信息层
    layers.append({
        "layer_name": "文档元信息层",
        "level": 1,
        "description": "文档的基础信息和格式设置",
        "content": {
            "file_size": analysis["file_info"]["size_bytes"],
            "paragraph_count": analysis["file_info"]["paragraph_count"],
            "tables_count": analysis["file_info"]["tables_count"],
            "sections_count": analysis["file_info"]["sections_count"],
            "styles_used": analysis["document_structure"]["styles_used"]
        }
    })
    
    # 层次2: 结构层次（标题层级）
    heading_hierarchy = {}
    for heading in analysis["content_analysis"]["headings"]:
        level = heading.get("level", 1)
        if level not in heading_hierarchy:
            heading_hierarchy[level] = []
        heading_hierarchy[level].append({
            "text": heading["text"],
            "style": heading["style"],
            "index": heading["index"]
        })
    
    layers.append({
        "layer_name": "结构层次（标题层级）",
        "level": 2,
        "description": "文档的标题层级结构，反映内容的组织方式",
        "content": heading_hierarchy
    })
    
    # 层次3: 内容语义层
    semantic_blocks = []
    for block in analysis["document_structure"]["hierarchical_structure"]:
        if block["heading"]:
            semantic_blocks.append({
                "heading": block["heading"]["text"],
                "heading_level": block["heading"].get("level", 1),
                "content_paragraphs": len(block["content"]),
                "content_preview": [p["text"][:50] for p in block["content"][:3] if p["text"]]
            })
    
    layers.append({
        "layer_name": "内容语义层",
        "level": 3,
        "description": "基于标题-内容块的内容组织",
        "content": semantic_blocks
    })
    
    # 层次4: 文本格式层
    format_analysis = {
        "bold_text": [],
        "italic_text": [],
        "colored_text": [],
        "different_fonts": set(),
        "different_sizes": set()
    }
    
    for para in analysis["document_structure"]["paragraphs"]:
        for run in para["runs"]:
            if run["bold"]:
                format_analysis["bold_text"].append(run["text"][:30])
            if run["italic"]:
                format_analysis["italic_text"].append(run["text"][:30])
            if run["font_color"]:
                format_analysis["colored_text"].append(run["text"][:30])
            if run["font_name"]:
                format_analysis["different_fonts"].add(run["font_name"])
            if run["font_size"]:
                format_analysis["different_sizes"].add(run["font_size"])
    
    format_analysis["different_fonts"] = list(format_analysis["different_fonts"])
    format_analysis["different_sizes"] = list(format_analysis["different_sizes"])
    
    layers.append({
        "layer_name": "文本格式层",
        "level": 4,
        "description": "文本的格式信息（加粗、斜体、颜色、字体等）",
        "content": format_analysis
    })
    
    # 层次5: 列表结构层
    list_structure = []
    for para in analysis["document_structure"]["paragraphs"]:
        if para["is_list"]:
            list_structure.append({
                "text": para["text"][:100],
                "style": para["style"],
                "index": para["index"]
            })
    
    layers.append({
        "layer_name": "列表结构层",
        "level": 5,
        "description": "文档中的列表项和项目符号",
        "content": list_structure
    })
    
    # 层次6: 表格数据层
    if analysis["content_analysis"]["tables"]:
        layers.append({
            "layer_name": "表格数据层",
            "level": 6,
            "description": "文档中的表格结构和数据",
            "content": analysis["content_analysis"]["tables"]
        })
    
    return layers

if __name__ == "__main__":
    docx_path = "Demo文档.docx"
    
    print("="*80)
    print("Demo文档.docx 深度拆解分析")
    print("="*80)
    print(f"\n正在分析: {docx_path}\n")
    
    # 分析文档结构
    analysis = analyze_docx_structure(docx_path)
    
    # 识别语义层次
    semantic_layers = identify_semantic_layers(analysis)
    analysis["semantic_layers"] = semantic_layers
    
    # 输出分析结果
    print("\n" + "="*80)
    print("【层次1】文档元信息层")
    print("="*80)
    print(json.dumps(analysis["semantic_layers"][0]["content"], indent=2, ensure_ascii=False))
    
    print("\n" + "="*80)
    print("【层次2】结构层次（标题层级）")
    print("="*80)
    print(json.dumps(analysis["semantic_layers"][1]["content"], indent=2, ensure_ascii=False))
    
    print("\n" + "="*80)
    print("【层次3】内容语义层")
    print("="*80)
    for i, block in enumerate(analysis["semantic_layers"][2]["content"], 1):
        print(f"\n块 {i}: {block['heading']} (级别{block['heading_level']})")
        print(f"  内容段落数: {block['content_paragraphs']}")
        if block['content_preview']:
            print(f"  内容预览:")
            for preview in block['content_preview']:
                print(f"    - {preview}...")
    
    print("\n" + "="*80)
    print("【层次4】文本格式层")
    print("="*80)
    format_info = analysis["semantic_layers"][3]["content"]
    print(f"使用的字体: {format_info['different_fonts']}")
    print(f"使用的字号: {format_info['different_sizes']}")
    print(f"加粗文本数量: {len(format_info['bold_text'])}")
    print(f"斜体文本数量: {len(format_info['italic_text'])}")
    print(f"彩色文本数量: {len(format_info['colored_text'])}")
    
    if analysis["semantic_layers"][4]["content"]:
        print("\n" + "="*80)
        print("【层次5】列表结构层")
        print("="*80)
        print(f"列表项数量: {len(analysis['semantic_layers'][4]['content'])}")
        for i, item in enumerate(analysis["semantic_layers"][4]["content"][:5], 1):
            print(f"  {i}. {item['text']}...")
    
    if len(analysis["semantic_layers"]) > 5:
        print("\n" + "="*80)
        print("【层次6】表格数据层")
        print("="*80)
        for i, table in enumerate(analysis["semantic_layers"][5]["content"], 1):
            print(f"\n表格 {i}:")
            print(f"  行数: {table['rows']}, 列数: {table['columns']}")
            if table['structure']:
                print(f"  表格内容预览:")
                for row in table['structure'][:3]:
                    print(f"    {row}")
    
    # 保存完整分析结果
    output_file = "demo_docx_analysis.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(analysis, f, indent=2, ensure_ascii=False)
    
    print(f"\n\n✅ 完整分析结果已保存到: {output_file}")
    print(f"   总层次数: {len(semantic_layers)}")
    print(f"   总段落数: {analysis['file_info']['paragraph_count']}")
    print(f"   总表格数: {analysis['file_info']['tables_count']}")

